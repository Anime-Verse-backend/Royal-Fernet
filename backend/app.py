
import os
import uuid
import json
import psycopg2
import psycopg2.extras
import io
import qrcode
import logging
from docx import Document
from docx.shared import Inches
from flask import send_file, send_from_directory, request, jsonify
from dotenv import load_dotenv
from flask import Flask
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime

# --- App Initialization ---
load_dotenv()
app = Flask(__name__)
logging.basicConfig(level=logging.INFO)

# --- CORS Configuration ---
FRONTEND_URL = os.getenv('FRONTEND_URL', 'http://localhost:9002')
CORS(app, resources={r"/*": {"origins": [FRONTEND_URL, "http://localhost:9002"], "supports_credentials": True}})

# --- Database Connection Helper ---
def get_db_connection():
    try:
        # Render provides the DATABASE_URL env var.
        conn_string = os.getenv('DATABASE_URL')
        if not conn_string:
            raise ValueError("DATABASE_URL environment variable is not set.")
        
        connection = psycopg2.connect(conn_string)
        app.logger.info("Successfully connected to the PostgreSQL database.")
        return connection
    except psycopg2.Error as e:
        app.logger.error(f"Error connecting to PostgreSQL database: {e}")
        return None

# --- File Upload Configuration ---
UPLOAD_FOLDER = os.environ.get('UPLOAD_FOLDER_PATH', os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads'))
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- Helper Functions ---
def make_image_url_absolute(path):
    if not path or path.startswith(('http://', 'https://')):
        return path
    api_base = os.getenv('API_BASE_URL', request.host_url)
    return f"{api_base.rstrip('/')}{path}"

def process_image_paths_for_response(image_data_json):
    if not image_data_json:
        return []
    try:
        # Data can be a list of dicts from psycopg2 or a JSON string from form
        image_data = image_data_json if isinstance(image_data_json, list) else json.loads(image_data_json)
        
        processed_data = []
        for item in image_data:
            if isinstance(item, dict) and 'imageUrl' in item:
                item['imageUrl'] = make_image_url_absolute(item['imageUrl'])
            processed_data.append(item)
        return processed_data
    except (json.JSONDecodeError, TypeError):
        app.logger.error(f"Could not parse or process image JSON: {image_data_json}")
        return []

def format_colombian_pesos(amount):
    try:
        return f"${int(amount):,}".replace(",", ".")
    except (ValueError, TypeError):
        return str(amount)

# --- API Routes ---

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/api/register', methods=['POST'])
def register_user():
    data = request.get_json()
    if not data or not data.get('name') or not data.get('email') or not data.get('password'):
        return jsonify({'error': 'Missing required fields'}), 400

    hashed_password = generate_password_hash(data['password'])
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            sql = "INSERT INTO users (name, email, password_hash, role) VALUES (%s, %s, %s, %s)"
            cursor.execute(sql, (data['name'], data['email'], hashed_password, 'user'))
        conn.commit()
        return jsonify({'message': 'User registered successfully'}), 201
    except psycopg2.IntegrityError:
        return jsonify({'error': 'Email already exists'}), 409
    finally:
        if conn:
            conn.close()

@app.route('/api/login', methods=['POST'])
def login_user():
    data = request.get_json()
    if not data or not data.get('email') or not data.get('password'):
        return jsonify({'error': 'Email and password are required'}), 400
    
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            sql = "SELECT * FROM users WHERE email = %s AND role = 'admin'"
            cursor.execute(sql, (data['email'],))
            user = cursor.fetchone()
            if user and check_password_hash(user['password_hash'], data['password']):
                return jsonify({'message': 'Login successful', 'user': {'name': user['name'], 'email': user['email']}}), 200
            else:
                return jsonify({'error': 'Invalid credentials or not an admin'}), 401
    finally:
        if conn:
            conn.close()

@app.route('/api/products', methods=['GET', 'POST'])
def handle_products():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        if request.method == 'POST':
            image_paths = []
            for i in range(1, 5):
                file_key = f'image{i}'
                url_key = f'imageUrl{i}'
                if file_key in request.files and request.files[file_key].filename != '':
                    file = request.files[file_key]
                    if file and allowed_file(file.filename):
                        filename = secure_filename(f"{uuid.uuid4()}_{file.filename}")
                        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                        image_paths.append(f"/uploads/{filename}")
                elif request.form.get(url_key):
                    image_paths.append(request.form.get(url_key))
            
            product_data = {
                'id': str(uuid.uuid4()),
                'name': request.form.get('name'),
                'description': request.form.get('description'),
                'category': request.form.get('category'),
                'price': request.form.get('price'),
                'discount': request.form.get('discount', 0),
                'stock': request.form.get('stock', 100),
                'images': json.dumps(image_paths),
                'is_featured': request.form.get('isFeatured') == 'on'
            }
            with conn.cursor() as cursor:
                sql = """INSERT INTO products (id, name, description, category, price, discount, stock, images, is_featured) 
                         VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                cursor.execute(sql, tuple(product_data.values()))
            conn.commit()
            product_data['images'] = process_image_paths_for_response(product_data['images'])
            return jsonify(product_data), 201
        
        # GET Products
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            query = request.args.get('q')
            if query:
                search_term = f"%{query}%"
                sql = "SELECT * FROM products WHERE name ILIKE %s OR category ILIKE %s ORDER BY created_at DESC"
                cursor.execute(sql, (search_term, search_term))
            else:
                sql = "SELECT * FROM products ORDER BY created_at DESC"
                cursor.execute(sql)
            
            products = [dict(row) for row in cursor.fetchall()]
            for p in products:
                # This is a simple list of paths, not dicts
                if p.get('images'):
                    p['images'] = [make_image_url_absolute(path) for path in p['images'] if path]
            return jsonify(products)
    finally:
        if conn:
            conn.close()

@app.route('/api/products/<string:product_id>', methods=['GET', 'PUT', 'DELETE'])
def handle_product(product_id):
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            if request.method == 'GET':
                cursor.execute("SELECT * FROM products WHERE id = %s", (product_id,))
                product = cursor.fetchone()
                if product:
                    product_dict = dict(product)
                    if product_dict.get('images'):
                        product_dict['images'] = [make_image_url_absolute(path) for path in product_dict['images'] if path]
                    return jsonify(product_dict)
                return jsonify({'error': 'Product not found'}), 404
            
            elif request.method == 'PUT':
                api_base_url = os.getenv('API_BASE_URL', request.host_url).rstrip('/')
                new_image_paths = []
                for i in range(1, 5):
                    if f'image{i}' in request.files and request.files[f'image{i}'].filename != '':
                        file = request.files[f'image{i}']
                        if file and allowed_file(file.filename):
                            filename = secure_filename(f"{uuid.uuid4()}_{file.filename}")
                            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                            new_image_paths.append(f"/uploads/{filename}")
                    elif request.form.get(f'imageUrl{i}'):
                        existing_url = request.form.get(f'imageUrl{i}')
                        if existing_url.startswith(api_base_url):
                            new_image_paths.append(existing_url.replace(api_base_url, '', 1))
                        else:
                            new_image_paths.append(existing_url)
                
                sql = """UPDATE products SET name=%s, description=%s, category=%s, price=%s, discount=%s, stock=%s, images=%s, is_featured=%s
                         WHERE id=%s"""
                values = (
                    request.form.get('name'), request.form.get('description'), request.form.get('category'),
                    request.form.get('price'), request.form.get('discount', 0), request.form.get('stock', 100),
                    json.dumps(new_image_paths), request.form.get('isFeatured') == 'on', product_id
                )
                cursor.execute(sql, values)
                conn.commit()
                return jsonify({'message': 'Product updated successfully'})

            elif request.method == 'DELETE':
                cursor.execute("DELETE FROM products WHERE id = %s", (product_id,))
                conn.commit()
                return jsonify({'message': 'Product deleted'})
    finally:
        if conn:
            conn.close()

@app.route('/api/admins', methods=['GET', 'POST'])
def handle_admins():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            if request.method == 'POST':
                data = request.get_json()
                if not data or not data.get('name') or not data.get('email') or not data.get('password'):
                    return jsonify({'error': 'Missing required fields'}), 400
                hashed_password = generate_password_hash(data['password'])
                sql = "INSERT INTO users (name, email, password_hash, role) VALUES (%s, %s, %s, 'admin')"
                cursor.execute(sql, (data['name'], data['email'], hashed_password))
                conn.commit()
                return jsonify({'name': data['name'], 'email': data['email']}), 201
            
            # GET Admins
            query = request.args.get('q')
            sql = "SELECT id, name, email FROM users WHERE role = 'admin'"
            if query:
                sql += " AND (name ILIKE %s OR email ILIKE %s)"
                cursor.execute(sql, (f"%{query}%", f"%{query}%"))
            else:
                cursor.execute(sql)
            admins = [dict(row) for row in cursor.fetchall()]
            return jsonify(admins)
    except psycopg2.IntegrityError:
        return jsonify({'error': 'Admin with that email already exists'}), 409
    finally:
        if conn:
            conn.close()

@app.route('/api/admins/<int:admin_id>', methods=['DELETE'])
def handle_admin(admin_id):
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            cursor.execute("SELECT COUNT(*) as count FROM users WHERE role = 'admin'")
            admin_count = cursor.fetchone()['count']
            if admin_count <= 1:
                return jsonify({'error': 'Cannot delete the last administrator'}), 400
            
            cursor.execute("DELETE FROM users WHERE id = %s AND role = 'admin'", (admin_id,))
            rows_affected = cursor.rowcount
            if rows_affected == 0:
                return jsonify({'error': 'Administrator not found'}), 404
            conn.commit()
            return jsonify({'message': 'Administrator deleted successfully'})
    finally:
        if conn:
            conn.close()

@app.route('/api/settings', methods=['GET', 'POST'])
def handle_settings():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            if request.method == 'POST':
                data = request.form
                api_base_url = os.getenv('API_BASE_URL', request.host_url).rstrip('/')
                
                hero_images_json = data.get('heroImages', '[]')
                hero_images_data = json.loads(hero_images_json)
                processed_hero_images = []

                for index, slide in enumerate(hero_images_data):
                    file_key = f'heroImageFile_{index}'
                    url_key = f'heroImageUrl_{index}'

                    # 1. Check for a newly uploaded file
                    if file_key in request.files and request.files[file_key].filename != '':
                        file = request.files[file_key]
                        if file and allowed_file(file.filename):
                            filename = secure_filename(f"setting_hero_{uuid.uuid4()}_{file.filename}")
                            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                            slide['imageUrl'] = f"/uploads/{filename}"
                    
                    # 2. Check for an existing URL from hidden inputs or the original JSON
                    elif data.get(url_key):
                        existing_url = data.get(url_key)
                        # Strip the base URL to store a relative path, but only if it's there
                        if existing_url and isinstance(existing_url, str) and existing_url.startswith(api_base_url):
                             slide['imageUrl'] = existing_url.replace(api_base_url, '', 1)
                        else:
                             slide['imageUrl'] = existing_url

                    processed_hero_images.append(slide)

                sql = """INSERT INTO settings (id, hero_images, featured_collection_title, featured_collection_description, promo_section_title, promo_section_description, promo_section_video_url, phone, contact_email, twitter_url, instagram_url, facebook_url)
                         VALUES (1, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                         ON CONFLICT (id) DO UPDATE SET
                         hero_images = EXCLUDED.hero_images,
                         featured_collection_title = EXCLUDED.featured_collection_title,
                         featured_collection_description = EXCLUDED.featured_collection_description,
                         promo_section_title = EXCLUDED.promo_section_title,
                         promo_section_description = EXCLUDED.promo_section_description,
                         promo_section_video_url = EXCLUDED.promo_section_video_url,
                         phone = EXCLUDED.phone,
                         contact_email = EXCLUDED.contact_email,
                         twitter_url = EXCLUDED.twitter_url,
                         instagram_url = EXCLUDED.instagram_url,
                         facebook_url = EXCLUDED.facebook_url
                         RETURNING *"""
                values = (
                    json.dumps(processed_hero_images), data.get('featuredCollectionTitle'), data.get('featuredCollectionDescription'),
                    data.get('promoSectionTitle'), data.get('promoSectionDescription'), data.get('promoSectionVideoUrl'),
                    data.get('phone'), data.get('contactEmail'), data.get('twitterUrl'), data.get('instagramUrl'), data.get('facebookUrl')
                )
                cursor.execute(sql, values)
                settings = dict(cursor.fetchone())
                conn.commit()
                if settings and settings.get('hero_images'):
                    settings['hero_images'] = process_image_paths_for_response(settings['hero_images'])
                return jsonify(settings)
            
            # GET request
            cursor.execute("SELECT * FROM settings WHERE id = 1")
            settings = cursor.fetchone()
            if settings:
                settings_dict = dict(settings)
                if settings_dict.get('hero_images'):
                    settings_dict['hero_images'] = process_image_paths_for_response(settings_dict['hero_images'])

                return jsonify(settings_dict)
            return jsonify({}), 404
    except Exception as e:
        app.logger.error(f"Error in handle_settings: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/notifications', methods=['POST'])
def create_notification():
    data = request.get_json()
    if not data or 'message' not in data:
        return jsonify({'error': 'Message is required'}), 400
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            sql = "INSERT INTO notifications (message, image_url, link_url) VALUES (%s, %s, %s)"
            cursor.execute(sql, (data['message'], data.get('imageUrl'), data.get('linkUrl')))
        conn.commit()
        return jsonify({'message': 'Notification created'}), 201
    finally:
        if conn:
            conn.close()

@app.route('/api/notifications/latest', methods=['GET'])
def get_latest_notification():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            cursor.execute("SELECT * FROM notifications ORDER BY created_at DESC LIMIT 1")
            notification = cursor.fetchone()
            if not notification:
                return jsonify({'error': 'No notifications found'}), 404
            notification_dict = dict(notification)
            if notification_dict.get('image_url'):
                notification_dict['image_url'] = make_image_url_absolute(notification_dict['image_url'])
            return jsonify(notification_dict)
    finally:
        if conn:
            conn.close()

@app.route('/api/stores', methods=['GET', 'POST'])
def handle_stores():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            if request.method == 'POST':
                data = request.form
                image_url = data.get('imageUrl', '')
                if 'imageFile' in request.files and request.files['imageFile'].filename != '':
                    file = request.files['imageFile']
                    if file and allowed_file(file.filename):
                        filename = secure_filename(f"store_{uuid.uuid4()}_{file.filename}")
                        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                        image_url = f"/uploads/{filename}"
                sql = """INSERT INTO store_locations (name, address, city, phone, hours, map_embed_url, image_url)
                         VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING *"""
                values = (data.get('name'), data.get('address'), data.get('city'), data.get('phone'), data.get('hours'), data.get('mapEmbedUrl'), image_url)
                cursor.execute(sql, values)
                new_store = dict(cursor.fetchone())
                conn.commit()
                return jsonify(new_store), 201

            # GET all stores
            cursor.execute("SELECT * FROM store_locations")
            stores = [dict(row) for row in cursor.fetchall()]
            for s in stores:
                s['image_url'] = make_image_url_absolute(s.get('image_url'))
            return jsonify(stores)
    finally:
        if conn:
            conn.close()

@app.route('/api/stores/<int:store_id>', methods=['PUT', 'DELETE'])
def handle_store(store_id):
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            if request.method == 'PUT':
                data = request.form
                cursor.execute("SELECT image_url FROM store_locations WHERE id = %s", (store_id,))
                current_image = cursor.fetchone()['image_url']
                image_url = data.get('imageUrl', current_image)
                if 'imageFile' in request.files and request.files['imageFile'].filename != '':
                    file = request.files['imageFile']
                    if file and allowed_file(file.filename):
                        filename = secure_filename(f"store_{uuid.uuid4()}_{file.filename}")
                        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                        image_url = f"/uploads/{filename}"
                sql = """UPDATE store_locations SET name=%s, address=%s, city=%s, phone=%s, hours=%s, map_embed_url=%s, image_url=%s
                         WHERE id=%s RETURNING *"""
                values = (data.get('name'), data.get('address'), data.get('city'), data.get('phone'), data.get('hours'), data.get('mapEmbedUrl'), image_url, store_id)
                cursor.execute(sql, values)
                updated_store = dict(cursor.fetchone())
                conn.commit()
                return jsonify(updated_store)

            elif request.method == 'DELETE':
                cursor.execute("DELETE FROM store_locations WHERE id = %s", (store_id,))
                conn.commit()
                return jsonify({'message': 'Store deleted successfully'})
    finally:
        if conn:
            conn.close()

@app.route('/api/generate-invoice-docx', methods=['POST'])
def generate_invoice_docx():
    data = request.get_json()
    if not data or not data.get('customerName') or not data.get('items'):
        return jsonify({'error': 'Missing required fields'}), 400
    
    customer_name = data['customerName']
    items = data['items']
    grand_total = 0
    invoice_items = []
    
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            for item_data in items:
                product_id = item_data.get('productId')
                quantity = int(item_data.get('quantity', 1))

                cursor.execute("SELECT * FROM products WHERE id = %s FOR UPDATE", (product_id,))
                product = cursor.fetchone()

                if not product:
                    raise ValueError(f"Producto con ID {product_id} no encontrado.")
                if product['stock'] < quantity:
                    raise ValueError(f"Stock insuficiente para '{product['name']}'. Disponible: {product['stock']}, Solicitado: {quantity}.")

                new_stock = product['stock'] - quantity
                cursor.execute("UPDATE products SET stock = %s WHERE id = %s", (new_stock, product_id))
                
                price = float(product['price'])
                discount_percent = int(product.get('discount', 0))
                discounted_price = price - (price * discount_percent / 100)
                subtotal = discounted_price * quantity
                grand_total += subtotal
                
                invoice_items.append({
                    "name": product['name'],
                    "quantity": quantity,
                    "unit_price": discounted_price,
                    "subtotal": subtotal
                })
        conn.commit()

    except (ValueError, psycopg2.Error) as e:
        conn.rollback()
        app.logger.error(f"Error processing invoice transaction: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

    try:
        invoice_number = f"INV-{int(datetime.now().timestamp())}"
        qr_url = "https://royal-fernet.vercel.app"
        qr_img = qrcode.make(qr_url)
        qr_io = io.BytesIO()
        qr_img.save(qr_io, 'PNG')
        qr_io.seek(0)

        document = Document()
        document.add_heading('Factura - Royal-Fernet', 0)
        p = document.add_paragraph()
        p.add_run('Vendedor: ').bold = True
        p.add_run('Royal-Fernet\n')
        p.add_run('Cliente: ').bold = True
        p.add_run(f'{customer_name}\n')
        p.add_run('Fecha: ').bold = True
        p.add_run(f'{datetime.now().strftime("%Y-%m-%d")}\n')
        p.add_run('Número de Factura: ').bold = True
        p.add_run(invoice_number)

        document.add_heading('Productos Comprados', level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Producto'
        hdr_cells[1].text = 'Cantidad'
        hdr_cells[2].text = 'Precio Unitario'
        hdr_cells[3].text = 'Subtotal'

        for item in invoice_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item['name']
            row_cells[1].text = str(item['quantity'])
            row_cells[2].text = format_colombian_pesos(item['unit_price'])
            row_cells[3].text = format_colombian_pesos(item['subtotal'])

        p_total = document.add_paragraph()
        p_total.alignment = 2 
        run_total_label = p_total.add_run('TOTAL A PAGAR: ')
        run_total_label.bold = True
        run_total_value = p_total.add_run(format_colombian_pesos(grand_total))
        run_total_value.bold = True

        document.add_heading('Detalles de la Tienda', level=1)
        document.add_paragraph('Escanea este código QR para visitar nuestra tienda online:')
        document.add_picture(qr_io, width=Inches(1.5))

        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'factura_{customer_name.replace(" ", "_")}.docx'
        )
    except Exception as e:
        app.logger.error(f"Error generating DOCX: {e}")
        return jsonify({'error': 'Failed to generate invoice document'}), 500

# --- Database Viewer Endpoints ---

@app.route('/api/db/tables', methods=['GET'])
def get_db_tables():
    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT table_name FROM information_schema.tables 
                WHERE table_schema = 'public' AND table_type = 'BASE TABLE'
            """)
            tables = [row[0] for row in cursor.fetchall()]
            return jsonify(tables)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/db/tables/<string:table_name>', methods=['GET'])
def get_table_content(table_name):
    # Basic sanitation to prevent obvious SQL injection.
    # In a real-world app, use a more robust validation method.
    if not table_name.isalnum() and '_' not in table_name:
        return jsonify({'error': 'Invalid table name'}), 400

    conn = get_db_connection()
    if not conn: return jsonify({'error': 'Database connection failed'}), 500
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.DictCursor) as cursor:
            # The table name is sanitized, but it's generally safer to use psycopg2's sql module for identifiers
            from psycopg2 import sql
            query = sql.SQL("SELECT * FROM {}").format(sql.Identifier(table_name))
            cursor.execute(query)
            
            # Fetch all rows and convert them to a list of dicts
            rows = [dict(row) for row in cursor.fetchall()]
            
            # Convert non-serializable types like datetime or decimal
            def json_converter(o):
                if isinstance(o, datetime):
                    return o.isoformat()
                if isinstance(o, (float, int)):
                     return o
                if isinstance(o, (list, dict)):
                    try:
                        return json.dumps(o, ensure_ascii=False)
                    except TypeError:
                        return str(o)
                return str(o)
            
            # Apply converter to each value in each row
            for row in rows:
                for key, value in row.items():
                    row[key] = json_converter(value)
                    
            return jsonify(rows)
            
    except Exception as e:
        return jsonify({'error': f"Could not fetch table '{table_name}': {str(e)}"}), 500
    finally:
        if conn:
            conn.close()


if __name__ == '__main__':
    # This block is for local development.
    # When deployed on Render with Gunicorn, Gunicorn will bind to the host and port.
    # Render provides the PORT environment variable.
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)


    